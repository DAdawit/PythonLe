from docx import Document

# Create a new Document
doc = Document()

# Add title
doc.add_heading('Personal Portfolio Website Content', level=1)

# Add content sections
doc.add_heading('Home', level=2)
doc.add_paragraph("Hi, I'm Dawit Diriba – Full-Stack Web Developer")
doc.add_paragraph("Transforming ideas into scalable, robust, and user-friendly web solutions.")
doc.add_paragraph("[View My Work] [Contact Me]")

doc.add_heading('About Me', level=2)
doc.add_paragraph("I am a passionate Full-Stack Web Developer with over 4 years of experience creating dynamic and scalable web applications. My expertise lies in designing user-friendly interfaces with cutting-edge front-end frameworks like React, Next.js, and Vue.js, seamlessly integrated with back-end technologies such as Laravel and Node.js. I’m committed to delivering exceptional web experiences through clean code, performance optimization, and secure architectures.\n\n"
                  "When I’m not coding, I enjoy exploring the latest web technologies and solving complex challenges to bring my clients’ visions to life.")

doc.add_heading('Skills', level=2)
doc.add_paragraph("Frontend: HTML5, CSS3, JavaScript, TypeScript, React, Next.js, Vue.js, TailwindCSS, Vuetify, Redux, Vuex.")
doc.add_paragraph("Backend: Node.js, Express.js, Laravel.")
doc.add_paragraph("Databases: MySQL, MongoDB, PostgreSQL.")
doc.add_paragraph("APIs: REST, GraphQL, OpenAPI.")
doc.add_paragraph("Tools & Platforms: Strapi CMS, Sanity CMS, Docker.")
doc.add_paragraph("Soft Skills: Collaboration, self-motivation, adaptability, and learning new technologies.")

doc.add_heading('Experience', level=2)
doc.add_paragraph("Fullstack Developer | PropheciusTechnologies (Jul 2023 – Present)\n"
                  "- Led a complete website redesign for Tadreeb Training and Consultancy using Next.js and Laravel.\n"
                  "- Integrated JWT-based authentication for secure user sessions.\n"
                  "- Technologies: Next.js, Laravel, TypeScript, TailwindCSS.\n")
doc.add_paragraph("Frontend Developer | Elihu Technologies (Nov 2022 – Jul 2023)\n"
                  "- Developed an e-commerce platform with real-time data updates using Next.js and GraphQL.\n"
                  "- Engineered a comprehensive admin panel for managing products, orders, and customers.\n"
                  "- Technologies: Next.js, GraphQL, TailwindCSS.\n")
doc.add_paragraph("Frontend Developer | Ashewa Technologies (Jun 2021 – Jan 2022)\n"
                  "- Designed dynamic web portals using Vue.js and Vuetify.\n"
                  "- Integrated advanced state management with Vuex.\n"
                  "- Technologies: Vue.js, Vuetify, GraphQL.\n")

doc.add_heading('Projects', level=2)
doc.add_paragraph("1. Tadreeb Training and Consultancy (2023)\n"
                  "- A modern, high-performing website for a Dubai-based training provider.\n"
                  "- Built with Next.js and Laravel, optimizing performance and SEO.\n"
                  "- Technologies: Next.js, TailwindCSS, Laravel.\n")
doc.add_paragraph("2. iSellSmartUSA (2023)\n"
                  "- A platform for trading used iPhones with dynamic content management.\n"
                  "- Implemented a headless CMS with Strapi for streamlined backend processes.\n"
                  "- Technologies: React, Vite, TailwindCSS.\n")
doc.add_paragraph("3. E-Commerce Platform (2022)\n"
                  "- A real-time e-commerce site with robust admin capabilities.\n"
                  "- Integrated Next.js with GraphQL for enhanced performance.\n"
                  "- Technologies: Next.js, GraphQL, TailwindCSS.\n")

doc.add_heading('Contact', level=2)
doc.add_paragraph("Email: dawitccnt@gmail.com")
doc.add_paragraph("Phone: +2519-3620-7512")
doc.add_paragraph("Location: Addis Ababa, Ethiopia")
doc.add_paragraph("LinkedIn/GitHub Links: Add personalized links")

# Save the document to a file
file_path = "./Dawit_Portfolio_Content.docx"
doc.save(file_path)

file_path
